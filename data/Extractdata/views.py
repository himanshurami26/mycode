from django.shortcuts import render
# from data.Extractdata.serializers import CountrySerializer
import docx
import os
import pandas as pd
from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.parsers import FileUploadParser,MultiPartParser
from rest_framework import views
from rest_framework.parsers import JSONParser 
from rest_framework.response import Response
from .models import ContryDetails,Category,StudentDetails,filedetails
from .serializers import CountrySerializer,CategorySerializer,StudentDetailsserializer,FiledetailsSerializer
import json
import re

# Stu. Dependent,Student,Visitor

# class DataView(APIView):
#     parser_classes = (MultiPartParser,)

#     def post(self, request, format=None):
#         file_obj = request.FILES['file'] 
#         ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/free lancing/' + str(file_obj)
#         Text = docx.Document(ROOT_DIR)
#         data = {}
#         res = {}
#         paragraphs = Text.paragraphs
#         heading = []
#         text = []
#         last_heading = None
#         for i in range(1, len(Text.paragraphs) - 1):
#             data[i] = Text.paragraphs[i].text.split()
#             dl = ' '.join(str(e) for e in data[i])
#             if len(dl) <= 5:
#                 continue
#             elif 5 < len(dl) <= 100:
#                 res[dl] = None
#                 last_heading = dl
#                 continue
#             else:
#                 res[last_heading] = dl
#                 continue
#         category_name = 'studentvisa'
#         user_obj =Category.objects.filter(id =2).delete()
#         #save the needed values 
#         # user_obj.save()
#         return_dict = {'status':True,'code':100}
#         for key, value in res.items():
#             heading.append(key)
#             text.append(value)
#             if key == None:
#                 key = 'Introduction'
#             if value == None:
#                 break
#             user = StudentDetails.objects.create(country_id_id=1,category_id_id=1,heading=key,peragraph=value)
#             # user = StudentDetails.objects.create(country_id=country_id,category_id=category_id,heading=x,peragraph=y)
#             user.save()
#         return Response(res)

class DataView(APIView):
    parser_classes = (MultiPartParser,)

    def get(self, request, format=None):
        for x in os.listdir('free lancing'):
            if x.endswith('.docx'):
                ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/free lancing/' + str(x)
                Text = docx.Document(ROOT_DIR)
                paragraphs = Text.paragraphs
                last_heading = None
                file_name = str(x)
                fix_final = []
                # print(file_name,'file_name')
                try:
                    qs = filedetails.objects.get(file=file_name)
                    serializer = FiledetailsSerializer(qs)
                    country = serializer.data
                    service_type_name = country['service_type_name']
                    # print(service_type_name,'service_type_name')
                    letter_type_name = country['letter_type_name']
                    # print(letter_type_name,'letter_type_name')
                    country_name = country['country_name']
                    # print(country_name,'country_name')
                    rop = []
                    data = {}
                    res = {}
                    if service_type_name == 'Stu. Dependent' or service_type_name == 'Student' or service_type_name == 'Visitor':
                        try:
                            category_data = Category.objects.get(category_name=letter_type_name)
                            category_serializer = CategorySerializer(category_data)
                            category = category_serializer.data
                            category_foreign_key = category['id']
                            # print(category_foreign_key,'category_foreign_key')
                            user = ContryDetails.objects.get(country_name=country_name)
                            contry_serializer = CountrySerializer(user)
                            country123 = contry_serializer.data
                            country_forign_key = country123['id']
                            # print(country_forign_key,'country_forign_key')
                            for i in range(1, len(paragraphs) - 1):
                                data[i] = paragraphs[i].text.split()
                                dl = ' '.join(str(e) for e in data[i])
                                if len(dl) <= 5:
                                    continue
                                elif 5 < len(dl) <= 100:
                                    res[dl] = None
                                    last_heading = dl
                                    continue
                                else:
                                    res[last_heading] = dl
                                    continue
                            for key, value in res.items():
                                if key == None:
                                    key = 'Introduction'
                                if value == None:
                                    break
                                try:
                                    user = StudentDetails.objects.create(country_id_id=country_forign_key,category_id_id=category_foreign_key,heading=key,peragraph=value)
                                    user.save()   
                                except:
                                    print("error")
                            # rop.append(res)
                        except:
                            continue
                    # output  =  fix_final.append(rop) 
                except:
                    continue  
            # json = {
            #     'final_response' : output
            # }
        return Response('success')


class expostdataindb(APIView):
    def post(self, request):
        country_name = 'Canada'
        user_obj =ContryDetails.objects.create(country_name=country_name)
        #save the needed values 
        user_obj.save()
        return_dict = {'status':True,'code':100}
        return HttpResponse(json.dumps(return_dict))





# "Requesting the concerned official to kindly do the needful and oblige.": null,
# "Thanking you in anticipation…": null,
# "________________": null,
# "SHREY B. PATEL": null
    

#REASONS TO CHOOSE UNITED KINGDOM
#Why Imagine Education Australia
#REASONS TO CHOOSE CANADA
#WHY AUSTRALIA
# REASONS TO SELECT CANADA AS STUDY DESTINATION
#REASONS NOT TO STUDY ELSEWHERE BUT USA
#REASONS TO SELECT CANADA AS STUDY DESTINATION

# print(rop,'rop')
# print(file_name)
# print(country,'country')
# tempfilename, tempfile_extension = os.path.splitext(file_name)
# if tempfile_extension == '.docx':
#     print(tempfilename, tempfile_extension)
# phoneNumRegex = re.compile(r'(\d\d\d)-(\d\d\d-\d\d\d\d)')
# mo = phoneNumRegex.search(file_name)
# print(mo.groups())

# print(country,'country')
# file_name = country['file']
    # user = StudentDetails.objects.create(country_id_id=country_id,category_id_id=category_id,heading=key,peragraph=value)
    # user.save()

# user_obj = Category.objects.create(category_name=category_name)
# user_obj.save()
# print(user_obj,'user')

# user_obj =ContryDetails.objects.create(country_name=country_name)
# #save the needed values 
# user_obj.save()
# user = ContryDetails.objects.get(country_name=country_name)

# user_obj = ContryDetails.objects.create(country_name=country_name)
# user_obj.save()
# print(user_obj,'user')
# country_name = 'Canada'
# # user_obj =ContryDetails.objects.create(country_name=country_name)
# # #save the needed values 
# # user_obj.save()
# user = ContryDetails.objects.get(country_name=country_name)
# print(user,'user')
# return_dict = {'status':True,'code':100}





